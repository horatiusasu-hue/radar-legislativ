#!/usr/bin/env python3
"""
Radar — supraveghere transparență decizională ANAF.

Ce face, la fiecare rulare:
  1. citește pagina de transparență decizională a ANAF
  2. compară cu ce e deja arhivat
  3. descarcă arhivele proiectelor noi
  4. le dezarhivează
  5. convertește documentele Word/Excel în text
  6. actualizează un index (index.json)
  7. afișează ce e nou

Rulare:  python3 radar_transparenta.py
Dependințe:  pip install requests beautifulsoup4
Opțional (pentru .doc vechi):  LibreOffice instalat
"""

import json
import re
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from urllib.parse import urljoin

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    sys.exit("Lipsesc dependințele. Rulează:  pip install requests beautifulsoup4")

# ─────────────────────────────────────────────────────────────
# CONFIGURARE
# ─────────────────────────────────────────────────────────────

PAGINA = "https://www.anaf.ro/anaf/internet/ANAF/transparenta_decizionala/"
ARHIVA = Path("arhiva")          # aici se salvează tot
INDEX = ARHIVA / "index.json"
TIMEOUT = 60

# Termenul legal de observații, în zile calendaristice (Legea 52/2003)
ZILE_OBSERVATII = 10

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36",
    "Accept-Language": "ro-RO,ro;q=0.9",
}

LUNI = {
    "ianuarie": 1, "februarie": 2, "martie": 3, "aprilie": 4,
    "mai": 5, "iunie": 6, "iulie": 7, "august": 8,
    "septembrie": 9, "octombrie": 10, "noiembrie": 11, "decembrie": 12,
}


# ─────────────────────────────────────────────────────────────
# INDEX
# ─────────────────────────────────────────────────────────────

def incarca_index() -> dict:
    if INDEX.exists():
        return json.loads(INDEX.read_text(encoding="utf-8"))
    return {"proiecte": {}}


def salveaza_index(index: dict) -> None:
    ARHIVA.mkdir(parents=True, exist_ok=True)
    INDEX.write_text(
        json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8"
    )


# ─────────────────────────────────────────────────────────────
# CITIRE PAGINĂ
# ─────────────────────────────────────────────────────────────

def parseaza_data(text: str):
    """'10 august 2026' -> date(2026, 8, 10)"""
    m = re.search(r"(\d{1,2})\s+([a-zăâîșț]+)\s+(\d{4})", text.lower())
    if not m:
        return None
    zi, luna, an = m.groups()
    if luna not in LUNI:
        return None
    return datetime(int(an), LUNI[luna], int(zi)).date()


def citeste_proiecte() -> list[dict]:
    """Întoarce lista proiectelor de pe pagină: data, titlu, url arhivă."""
    r = requests.get(PAGINA, headers=HEADERS, timeout=TIMEOUT)
    r.raise_for_status()
    sup = BeautifulSoup(r.text, "html.parser")

    proiecte = []
    for a in sup.find_all("a", href=True):
        href = urljoin(PAGINA, a["href"])
        if not href.lower().endswith(".zip"):
            continue

        # contextul din jurul linkului conține data și titlul
        bloc = a.find_parent(["li", "div", "tr", "p"]) or a.parent
        context = bloc.get_text(" ", strip=True) if bloc else ""
        data = parseaza_data(context)

        titlu = re.sub(r"\s*\d{1,2}\s+\w+\s+\d{4}\s*\|?\s*", "", context)
        titlu = titlu.replace("Detalii proiect", "").strip(" |–- ")

        proiecte.append({
            "url": href,
            "data": data.isoformat() if data else None,
            "titlu": titlu[:400],
            "id": href.rsplit("/", 1)[-1].replace(".zip", ""),
        })

    # deduplicare, păstrând ordinea
    vazute, unice = set(), []
    for p in proiecte:
        if p["id"] not in vazute:
            vazute.add(p["id"])
            unice.append(p)
    return unice


# ─────────────────────────────────────────────────────────────
# DESCĂRCARE ȘI EXTRAGERE
# ─────────────────────────────────────────────────────────────

def descarca(url: str, destinatie: Path) -> Path:
    destinatie.parent.mkdir(parents=True, exist_ok=True)
    with requests.get(url, headers=HEADERS, timeout=TIMEOUT, stream=True) as r:
        r.raise_for_status()
        with open(destinatie, "wb") as f:
            for bucata in r.iter_content(65536):
                f.write(bucata)
    return destinatie


def dezarhiveaza(zip_path: Path, unde: Path) -> list[Path]:
    unde.mkdir(parents=True, exist_ok=True)
    fisiere = []
    with zipfile.ZipFile(zip_path) as z:
        for nume in z.namelist():
            if nume.endswith("/"):
                continue
            # nume sigur, fără cale absolută sau ../
            sigur = Path(nume).name
            tinta = unde / sigur
            with z.open(nume) as sursa, open(tinta, "wb") as dest:
                dest.write(sursa.read())
            fisiere.append(tinta)
    return fisiere


def in_text(fisier: Path) -> str | None:
    """Convertește un document în text. Întoarce None dacă nu se poate."""
    ext = fisier.suffix.lower()

    if ext == ".docx":
        try:
            import docx
            d = docx.Document(str(fisier))
            parti = [p.text for p in d.paragraphs]
            for t in d.tables:
                for rand in t.rows:
                    parti.append(" | ".join(c.text.strip() for c in rand.cells))
            return "\n".join(parti)
        except Exception:
            pass

    if ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(fisier), data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"--- {ws.title} ---")
                for rand in ws.iter_rows(values_only=True):
                    if any(c is not None for c in rand):
                        out.append(" | ".join("" if c is None else str(c) for c in rand))
            return "\n".join(out)
        except Exception:
            pass

    if ext == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(str(fisier))
            out = []
            for ws in wb.sheets():
                out.append(f"--- {ws.name} ---")
                for i in range(ws.nrows):
                    out.append(" | ".join(str(v) for v in ws.row_values(i)))
            return "\n".join(out)
        except Exception:
            pass

    # .doc vechi și orice altceva: prin LibreOffice
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "txt:Text",
             "--outdir", str(fisier.parent), str(fisier)],
            check=True, capture_output=True, timeout=120,
        )
        txt = fisier.with_suffix(".txt")
        if txt.exists():
            return txt.read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass

    return None


# ─────────────────────────────────────────────────────────────
# PRINCIPAL
# ─────────────────────────────────────────────────────────────

def main() -> None:
    index = incarca_index()
    cunoscute = index["proiecte"]

    print(f"Citesc {PAGINA} …")
    try:
        proiecte = citeste_proiecte()
    except Exception as e:
        sys.exit(f"Nu am putut citi pagina: {e}\n"
                 f"Dacă pagina e randată cu JavaScript, vezi nota din README.")

    if not proiecte:
        print("ATENȚIE: n-am găsit niciun link .zip. Structura paginii s-a schimbat?")
        return

    print(f"Găsite pe pagină: {len(proiecte)} proiecte.\n")
    noi = []

    for p in proiecte:
        if p["id"] in cunoscute:
            continue

        print(f"NOU  {p['data'] or '?'}  {p['titlu'][:90]}")
        folder = ARHIVA / (p["data"] or "fara-data") / p["id"]

        try:
            zp = descarca(p["url"], folder / f"{p['id']}.zip")
            fisiere = dezarhiveaza(zp, folder / "continut")
        except Exception as e:
            print(f"     ! eroare la descărcare/dezarhivare: {e}")
            continue

        texte = folder / "text"
        texte.mkdir(exist_ok=True)
        convertite = []
        for f in fisiere:
            t = in_text(f)
            if t:
                (texte / (f.stem + ".txt")).write_text(t, encoding="utf-8")
                convertite.append(f.name)

        # termenul de observații
        termen = None
        if p["data"]:
            d = datetime.fromisoformat(p["data"]).date()
            termen = (d.toordinal() + ZILE_OBSERVATII)
            termen = datetime.fromordinal(termen).date().isoformat()

        cunoscute[p["id"]] = {
            **p,
            "termen_observatii": termen,
            "fisiere": [f.name for f in fisiere],
            "convertite": convertite,
            "arhivat_la": datetime.now(timezone.utc).isoformat(timespec="seconds"),
            "stare": "in_consultare",   # → "adoptat" / "abandonat", completat manual
            "act_publicat": None,        # → nr. M.Of., când apare
        }
        noi.append(cunoscute[p["id"]])
        print(f"     {len(fisiere)} fișiere, {len(convertite)} convertite în text")

    salveaza_index(index)

    print("\n" + "=" * 60)
    if not noi:
        print("Nimic nou. Nu se trimite nimic.")
        return

    azi = datetime.now().date()
    print(f"{len(noi)} proiect(e) nou(i):\n")
    for p in noi:
        print(f"  {p['titlu'][:100]}")
        print(f"    publicat: {p['data']}")
        if p["termen_observatii"]:
            t = datetime.fromisoformat(p["termen_observatii"]).date()
            zile = (t - azi).days
            stare = f"{zile} zile rămase" if zile >= 0 else f"EXPIRAT de {-zile} zile"
            print(f"    termen observații: {p['termen_observatii']}  ({stare})")
        print(f"    text extras în: arhiva/{p['data']}/{p['id']}/text/\n")


if __name__ == "__main__":
    main()
