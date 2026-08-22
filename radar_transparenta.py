#!/usr/bin/env python3
"""
Radar — supraveghere transparență decizională, mai multe ministere.

Ce face, la fiecare rulare:
  1. citește paginile de transparență decizională ale instituțiilor urmărite
  2. compară cu ce e deja arhivat
  3. descarcă arhivele/documentele proiectelor noi
  4. le dezarhivează și le convertește în text
  5. le încadrează pe domenii (cele 20 din Radar)
  6. caută dacă vreun proiect vechi a devenit act publicat  ← NOU
  7. actualizează index.json și afișează ce e nou

Rulare:  python3 radar_transparenta.py
Dependințe:  pip install requests beautifulsoup4
Opțional (pentru .doc vechi):  LibreOffice instalat
"""

import json
import re
import subprocess
import sys
import zipfile
from datetime import datetime, timezone
from difflib import SequenceMatcher
from pathlib import Path
from urllib.parse import urljoin

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    sys.exit("Lipsesc dependințele. Rulează:  pip install requests beautifulsoup4")


# ─────────────────────────────────────────────────────────────
# SURSE — proiecte în transparență decizională
# ─────────────────────────────────────────────────────────────

# instituție, adresă, zile termen observații, domenii implicite, confirmată
#
# „confirmată" = am verificat că adresa răspunde și întoarce proiecte reale.
# Sursele NEconfirmate rămân în listă intenționat, ca să apară zilnic în raport
# la „DE REPARAT". Nu le ștergem — o sursă ștearsă e o gaură pe care o uiți.
#
# Diferența contează în raport:
#   sursă confirmată care pică  = ALARMĂ  (s-a stricat ceva; poate ai ratat acte)
#   sursă neconfirmată care pică = TODO   (n-a mers niciodată; de găsit adresa)

SURSE = [
    ("ANAF", "https://www.anaf.ro/anaf/internet/ANAF/transparenta_decizionala/", 10, ["1", "3"], True),
    ("Ministerul Finanțelor", "https://mfinante.gov.ro/acasa/transparenta/proiecte-acte-normative", 10, ["1", "3"], True),
    ("Ministerul Economiei", "https://economie.gov.ro/proiecte-de-acte-normative-aflate-in-consultare-publica/", 30, ["11"], True),
    ("Ministerul Muncii", "https://mmuncii.gov.ro/transparenta-decizionala/", 30, ["2"], True),

    # ── neconfirmate: adresa veche a murit, cea nouă nu e încă găsită ──
    ("Ministerul Transporturilor", "https://www.mt.ro/web14/transparenta-decizionala/consultare-publica/acte-normative-in-avizare", 30, ["4"], False),
    ("Vama (AVR)", "https://www.customs.ro/info-publice/transparenta-decizionala", 10, ["14", "8"], False),
    ("Ministerul Mediului", "https://www.mmediu.ro/categorie/transparenta-decizionala/1", 30, ["15"], False),
    ("Ministerul Agriculturii", "https://www.madr.ro/transparenta-decizionala.html", 30, ["9"], False),

    # plasă de siguranță: agregatorul guvernamental.
    # ATENȚIE: se actualizează SĂPTĂMÂNAL și manual — e în urmă, nu în față.
    ("e-consultare (agregator)", "https://e-consultare.gov.ro/Consultare-public%C4%83", 10, [], False),
]

# ─────────────────────────────────────────────────────────────
# SURSE — acte deja publicate (pentru potrivirea proiect → act)
# ─────────────────────────────────────────────────────────────

SURSE_ACTE_PUBLICATE = [
    ("ANAF — alte acte normative",
     "https://www.anaf.ro/anaf/internet/ANAF/asistenta_contribuabili/legislatie/alte_acte_normative/"),
    ("Monitorul Oficial — sumar",
     "https://monitoruloficial.ro/"),
]


# ─────────────────────────────────────────────────────────────
# DOMENII — încadrare pe cuvinte-cheie
# ─────────────────────────────────────────────────────────────

DOMENII = {
    "1":  ("Fiscal general",        ["cod fiscal", "tva", "impozit pe profit", "microîntreprinder",
                                      "cod de procedură fiscală", "taxa pe valoarea adăugată",
                                      "rambursare", "decont", "inspecție fiscală", "antifraud"]),
    "2":  ("Salarizare și muncă",   ["contribuți", "salari", "contract individual de muncă", "revisal",
                                      "inspecția muncii", "securitate în muncă", "concediu", "pensi"]),
    "3":  ("Raportări și declarații", ["declaraț", "formular", "saf-t", "d406", "e-factura", "e-transport",
                                      "e-tva", "raportare", "spv", "e-case de marcat"]),
    "4":  ("Transport rutier",      ["transport rutier", "a.d.r.", "adr", "licență de transport",
                                      "tahograf", "mărfuri periculoase", "vehicul"]),
    "5":  ("Comerț cu amănuntul",   ["casă de marcat", "aparate de marcat", "bon fiscal",
                                      "protecția consumator", "etichetare"]),
    "6":  ("Construcții",           ["construcț", "autorizație de construire", "recepți",
                                      "inspectoratul de stat în construcții", "urbanism"]),
    "7":  ("HoReCa",                ["alimentație publică", "restaurant", "turism gastronomic",
                                      "unități de alimentație"]),
    "8":  ("Accize și carburanți",  ["acciz", "carburant", "motorin", "benzin", "antrepozit fiscal",
                                      "produse energetice", "alcool", "tutun"]),
    "9":  ("Agricultură",           ["agricol", "subvenț", "apia", "fermier", "cereale", "zootehni"]),
    "10": ("Imobiliare",            ["imobil", "locuinț", "cadastru", "carte funciar", "teren"]),
    "11": ("Producție și industrie",["iscir", "sudor", "prescripți tehnic", "instalații sub presiune",
                                      "autorizare tehnic", "industrie"]),
    "12": ("Farma și medical",      ["medicament", "farmac", "dispozitiv medical", "sanitar"]),
    "13": ("IT și servicii",        ["software", "digitalizare", "servicii informatic", "cloud"]),
    "14": ("Import-export",         ["vam", "antidumping", "taric", "import", "export", "tarif vamal"]),
    "15": ("Deșeuri și mediu",      ["deșeu", "mediu", "ambalaj", "reciclare", "emisii", "poluare"]),
    "16": ("Energie și utilități",  ["anre", "energie electric", "gaze natural", "furnizare energie"]),
    "17": ("Turism",                ["agenți de turism", "structuri de primire", "voucher de vacanț"]),
    "18": ("Finanțări, ajutor de stat", ["ajutor de stat", "schemă de finanțare", "fond", "minimis",
                                      "apel de proiecte", "grant"]),
    "19": ("Comerț cu ridicata",    ["taxare inversă", "comerț cu ridicata", "distribuți"]),
}

IMPLICITE = {"1", "2", "3"}


# ─────────────────────────────────────────────────────────────
# CONFIGURARE
# ─────────────────────────────────────────────────────────────

ARHIVA = Path("arhiva")
INDEX = ARHIVA / "index.json"
TIMEOUT = 60
EXTENSII = (".zip", ".pdf", ".doc", ".docx", ".rtf")

# ── Praguri de siguranță ─────────────────────────────────────
# NU sunt filtre de mărime. Un act normativ mare (ordonanță-trenuleț cu 40 de
# anexe) trece fără probleme — cea mai mare arhivă ANAF de până acum are 142 KB.
# Pragurile astea există doar contra unei arhive-bombă: un .zip mic care,
# desfăcut, umflă zeci de gigaocteți și blochează rularea.
#
# REGULA: nimic nu dispare în tăcere. Ce depășește pragul se PĂSTREAZĂ și se
# RAPORTEAZĂ, ca să te uiți tu. Un fals negativ tăcut e mai rău decât o alarmă.
MAX_ARHIVA_DESFACUTA = 400 * 1024 * 1024   # 400 MB desfăcut, total
MAX_FISIER_CONVERSIE = 60 * 1024 * 1024    # peste atât: se arhivează, nu se convertește
MAX_RAPORT_COMPRESIE = 200                 # zip de 1 MB care dă 200 MB = suspect

AVERTISMENTE: list[str] = []               # se afișează la final, vizibil


def avertizeaza(mesaj: str) -> None:
    """Zgomotos și inofensiv, nu tăcut și periculos."""
    AVERTISMENTE.append(mesaj)
    print(f"    !! ATENȚIE: {mesaj}")

HEADERS = {
    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0 Safari/537.36",
    "Accept-Language": "ro-RO,ro;q=0.9",
}

LUNI = {"ianuarie": 1, "februarie": 2, "martie": 3, "aprilie": 4, "mai": 5, "iunie": 6,
        "iulie": 7, "august": 8, "septembrie": 9, "octombrie": 10, "noiembrie": 11, "decembrie": 12}

# „Ordinul ... nr. 352/2022", „O.p.A.N.A.F. nr. 1757/2019", „HG nr. 1175/2007"
TIPAR_ACT = re.compile(r"nr\.?\s*(\d{1,5})\s*/\s*(\d{4})", re.IGNORECASE)


# ─────────────────────────────────────────────────────────────
# UTILITARE
# ─────────────────────────────────────────────────────────────

def incarca_index() -> dict:
    if INDEX.exists():
        d = json.loads(INDEX.read_text(encoding="utf-8"))
        d.setdefault("proiecte", {})
        d.setdefault("acte_vazute", [])
        return d
    return {"proiecte": {}, "acte_vazute": []}


def salveaza_index(index: dict) -> None:
    ARHIVA.mkdir(parents=True, exist_ok=True)
    INDEX.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")


def parseaza_data(text: str):
    t = text.lower()
    m = re.search(r"(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})", t)
    if m:
        zi, luna, an = (int(x) for x in m.groups())
        try:
            return datetime(an, luna, zi).date()
        except ValueError:
            pass
    m = re.search(r"(\d{1,2})\s+([a-zăâîșțş]+)\s+(\d{4})", t)
    if m:
        zi, luna, an = m.groups()
        if luna in LUNI:
            try:
                return datetime(int(an), LUNI[luna], int(zi)).date()
            except ValueError:
                pass
    return None


def acte_mentionate(text: str) -> set:
    """{'352/2022', '1757/2019'} — actele la care trimite un titlu."""
    return {f"{a}/{b}" for a, b in TIPAR_ACT.findall(text or "")}


def normalizeaza(text: str) -> str:
    t = (text or "").lower()
    for a, b in [("ă", "a"), ("â", "a"), ("î", "i"), ("ș", "s"), ("ş", "s"), ("ț", "t"), ("ţ", "t")]:
        t = t.replace(a, b)
    return re.sub(r"[^a-z0-9 ]+", " ", t)


def asemanare(a: str, b: str) -> float:
    return SequenceMatcher(None, normalizeaza(a), normalizeaza(b)).ratio()


def incadreaza(titlu: str, implicite: list) -> list:
    """Domeniile în care intră un proiect, după cuvinte-cheie din titlu."""
    t = normalizeaza(titlu)
    gasite = set(implicite)
    for cod, (_, cuvinte) in DOMENII.items():
        if any(normalizeaza(c) in t for c in cuvinte):
            gasite.add(cod)
    return sorted(gasite, key=lambda x: int(x))


# ─────────────────────────────────────────────────────────────
# CITIRE PAGINI DE TRANSPARENȚĂ
# ─────────────────────────────────────────────────────────────

def citeste_sursa(nume: str, url: str, zile: int, domenii_implicite: list) -> list[dict]:
    r = requests.get(url, headers=HEADERS, timeout=TIMEOUT)
    r.raise_for_status()
    sup = BeautifulSoup(r.text, "html.parser")

    gasite, vazute = [], set()
    for a in sup.find_all("a", href=True):
        href = urljoin(url, a["href"])
        if not href.lower().split("?")[0].endswith(EXTENSII):
            continue
        if href in vazute:
            continue
        vazute.add(href)

        bloc = a.find_parent(["li", "tr", "div", "p", "article"]) or a.parent
        context = bloc.get_text(" ", strip=True) if bloc else a.get_text(strip=True)
        data = parseaza_data(context)

        titlu = re.sub(r"\s*\d{1,2}[.\-/ ]\w+[.\-/ ]\d{4}\s*\|?\s*", " ", context)
        for gunoi in ("Detalii proiect", "Descarcă", "Download", "citeste mai mult", "CITEŞTE MAI MULT"):
            titlu = titlu.replace(gunoi, "")
        titlu = " ".join(titlu.split())[:400]

        if len(titlu) < 25:          # linkuri de navigare, nu proiecte
            continue

        gasite.append({
            "institutie": nume,
            "url": href,
            "data": data.isoformat() if data else None,
            "titlu": titlu,
            "id": re.sub(r"[^A-Za-z0-9._-]", "_", href.rsplit("/", 1)[-1])[:120],
            "zile_observatii": zile,
            "domenii": incadreaza(titlu, domenii_implicite),
            "acte_modificate": sorted(acte_mentionate(titlu)),
        })
    return gasite


# ─────────────────────────────────────────────────────────────
# DESCĂRCARE ȘI EXTRAGERE
# ─────────────────────────────────────────────────────────────

def descarca(url: str, destinatie: Path) -> Path:
    destinatie.parent.mkdir(parents=True, exist_ok=True)
    with requests.get(url, headers=HEADERS, timeout=TIMEOUT, stream=True) as r:
        r.raise_for_status()
        with open(destinatie, "wb") as f:
            for bucata in r.iter_content(65536):
                f.write(bucata)
    return destinatie


def desface(cale: Path, unde: Path) -> list[Path]:
    """Dezarhivează, dacă e arhivă. Altfel întoarce fișierul ca atare.

    Verifică ÎNAINTE de desfacere cât declară arhiva că ocupă desfăcută.
    Un act normativ real, oricât de mare, trece. O bombă e oprită și semnalată.
    """
    unde.mkdir(parents=True, exist_ok=True)
    if cale.suffix.lower() != ".zip":
        tinta = unde / cale.name
        tinta.write_bytes(cale.read_bytes())
        return [tinta]

    with zipfile.ZipFile(cale) as z:
        intrari = [i for i in z.infolist() if not i.is_dir()]
        desfacut = sum(i.file_size for i in intrari)
        comprimat = max(sum(i.compress_size for i in intrari), 1)
        raport = desfacut / comprimat

        if desfacut > MAX_ARHIVA_DESFACUTA or raport > MAX_RAPORT_COMPRESIE:
            avertizeaza(
                f"arhiva {cale.name} declară {desfacut / 1048576:.0f} MB desfăcuți "
                f"(raport de compresie {raport:.0f}:1). NU am desfăcut-o. "
                f"Arhiva brută e păstrată la {cale}. Verific-o manual."
            )
            return []                              # nu se desface, dar nici nu se șterge

        fisiere = []
        for i in intrari:
            tinta = unde / Path(i.filename).name   # nume sigur, fără ../
            with z.open(i) as sursa, open(tinta, "wb") as dest:
                dest.write(sursa.read())
            fisiere.append(tinta)
    return fisiere


def in_text(fisier: Path) -> str | None:
    ext = fisier.suffix.lower()

    if fisier.stat().st_size > MAX_FISIER_CONVERSIE:
        avertizeaza(
            f"{fisier.name} are {fisier.stat().st_size / 1048576:.0f} MB — "
            f"e arhivat, dar nu l-am convertit în text (ar bloca rularea). "
            f"Dacă e un act care te interesează, deschide-l direct."
        )
        return None

    if ext == ".docx":
        try:
            import docx
            d = docx.Document(str(fisier))
            parti = [p.text for p in d.paragraphs]
            for t in d.tables:
                for rand in t.rows:
                    parti.append(" | ".join(c.text.strip() for c in rand.cells))
            return "\n".join(parti)
        except Exception:
            pass

    if ext == ".xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(fisier), data_only=True)
            out = []
            for ws in wb.worksheets:
                out.append(f"--- {ws.title} ---")
                for rand in ws.iter_rows(values_only=True):
                    if any(c is not None for c in rand):
                        out.append(" | ".join("" if c is None else str(c) for c in rand))
            return "\n".join(out)
        except Exception:
            pass

    if ext == ".xls":
        try:
            import xlrd
            wb = xlrd.open_workbook(str(fisier))
            out = []
            for ws in wb.sheets():
                out.append(f"--- {ws.name} ---")
                for i in range(ws.nrows):
                    out.append(" | ".join(str(v) for v in ws.row_values(i)))
            return "\n".join(out)
        except Exception:
            pass

    try:  # .doc vechi, .pdf, .rtf — prin LibreOffice
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "txt:Text",
             "--outdir", str(fisier.parent), str(fisier)],
            check=True, capture_output=True, timeout=180,
        )
        txt = fisier.with_suffix(".txt")
        if txt.exists():
            return txt.read_text(encoding="utf-8", errors="replace")
    except Exception:
        pass

    return None


# ─────────────────────────────────────────────────────────────
# POTRIVIREA PROIECT → ACT PUBLICAT
# ─────────────────────────────────────────────────────────────

def citeste_acte_publicate() -> list[dict]:
    """Titluri de acte publicate recent, din sursele configurate."""
    acte = []
    for nume, url in SURSE_ACTE_PUBLICATE:
        try:
            r = requests.get(url, headers=HEADERS, timeout=TIMEOUT)
            r.raise_for_status()
            sup = BeautifulSoup(r.text, "html.parser")
            for el in sup.find_all(["li", "tr", "p", "h2", "h3", "a"]):
                text = " ".join(el.get_text(" ", strip=True).split())
                if len(text) < 40 or len(text) > 600:
                    continue
                if not re.search(r"ordin|hot[ăa]r[âa]re|lege|ordonan[țt]", text, re.IGNORECASE):
                    continue
                acte.append({"sursa": nume, "titlu": text[:400],
                             "acte": sorted(acte_mentionate(text)),
                             "data": (parseaza_data(text) or "").__str__() or None})
        except Exception as e:
            print(f"  ! {nume}: {e}")
    return acte


def cauta_potriviri(index: dict) -> list[dict]:
    """Pentru fiecare proiect încă „in_consultare", caută actul publicat."""
    deschise = [p for p in index["proiecte"].values() if p.get("stare") == "in_consultare"]
    if not deschise:
        return []

    print("\nCaut acte publicate care să corespundă proiectelor deschise…")
    acte = citeste_acte_publicate()
    if not acte:
        print("  (nicio sursă de acte publicate n-a răspuns)")
        return []
    print(f"  {len(acte)} titluri de acte citite.")

    potriviri = []
    for p in deschise:
        candidati = []
        for act in acte:
            scor, motiv = 0.0, ""

            # TREAPTA 1 — ancoră exactă: același act modificat
            comune = set(p.get("acte_modificate", [])) & set(act["acte"])
            if comune:
                scor = 0.90
                motiv = f"modifică același act: {', '.join(sorted(comune))}"

            # TREAPTA 2 — asemănare de titlu
            s = asemanare(p["titlu"], act["titlu"])
            if s > scor:
                scor, motiv = s, f"titluri asemănătoare ({s:.0%})"
            elif comune and s > 0.5:
                scor = min(0.98, scor + 0.05)
                motiv += f" + titluri asemănătoare ({s:.0%})"

            if scor >= 0.62:
                candidati.append({"scor": round(scor, 2), "motiv": motiv, **act})

        if candidati:
            candidati.sort(key=lambda c: -c["scor"])
            potriviri.append({"proiect": p, "candidati": candidati[:3]})
    return potriviri


# ─────────────────────────────────────────────────────────────
# PRINCIPAL
# ─────────────────────────────────────────────────────────────

def main() -> None:
    index = incarca_index()
    cunoscute = index["proiecte"]
    noi = []

    alarme, de_reparat = [], []

    print("SURSE DE PROIECTE\n" + "─" * 60)
    for nume, url, zile, dom, confirmata in SURSE:
        try:
            gasite = citeste_sursa(nume, url, zile, dom)
            print(f"{nume:<32} {len(gasite):>3} proiecte")
            if not gasite:
                # zero proiecte nu e „e liniște". E ori chiar liniște, ori
                # s-a schimbat structura paginii și nu mai vedem nimic.
                (alarme if confirmata else de_reparat).append(
                    f"{nume}: pagina răspunde, dar n-am găsit niciun document. "
                    f"Ori chiar nu e nimic în consultare, ori s-a schimbat structura paginii."
                )
        except Exception as e:
            print(f"{nume:<32}  EROARE: {str(e)[:70]}")
            (alarme if confirmata else de_reparat).append(f"{nume}: {str(e)[:110]}\n      {url}")
            continue

        for p in gasite:
            cheie = f"{nume}::{p['id']}"
            if cheie in cunoscute:
                continue

            folder = ARHIVA / (p["data"] or "fara-data") / re.sub(r"[^A-Za-z0-9._-]", "_", nume) / p["id"]
            try:
                fis = descarca(p["url"], folder / p["url"].rsplit("/", 1)[-1][:120])
                fisiere = desface(fis, folder / "continut")
            except Exception as e:
                print(f"    ! {p['titlu'][:60]} — {str(e)[:60]}")
                continue

            texte = folder / "text"
            texte.mkdir(exist_ok=True)
            convertite = []
            for f in fisiere:
                t = in_text(f)
                if t:
                    (texte / (f.stem + ".txt")).write_text(t, encoding="utf-8")
                    convertite.append(f.name)

            termen = None
            if p["data"]:
                d = datetime.fromisoformat(p["data"]).date()
                termen = datetime.fromordinal(d.toordinal() + p["zile_observatii"]).date().isoformat()

            cunoscute[cheie] = {
                **p,
                "termen_observatii": termen,
                "fisiere": [f.name for f in fisiere],
                "convertite": convertite,
                "arhivat_la": datetime.now(timezone.utc).isoformat(timespec="seconds"),
                "stare": "in_consultare",
                "act_publicat": None,
                "cale": str(folder).replace("\\", "/"),
            }
            noi.append(cunoscute[cheie])
            print(f"    NOU  {p['titlu'][:70]}  ({len(convertite)}/{len(fisiere)} în text)")

    potriviri = cauta_potriviri(index)
    salveaza_index(index)

    # ── RAPORT ──────────────────────────────────────────────
    azi = datetime.now().date()
    print("\n" + "═" * 60)

    if noi:
        print(f"\nPROIECTE NOI — {len(noi)}\n")
        for p in noi:
            dom = ", ".join(DOMENII[d][0] for d in p["domenii"] if d in DOMENII) or "neîncadrat"
            print(f"  {p['institutie']} · {p['data'] or '?'}")
            print(f"  {p['titlu'][:100]}")
            print(f"  domenii: {dom}")
            if p["termen_observatii"]:
                t = datetime.fromisoformat(p["termen_observatii"]).date()
                z = (t - azi).days
                print(f"  termen observații: {p['termen_observatii']}  "
                      f"({z} zile rămase)" if z >= 0 else
                      f"  termen observații: {p['termen_observatii']}  (EXPIRAT de {-z} zile)")
            print(f"  text: {p['cale']}/text/\n")
    else:
        print("\nNiciun proiect nou.\n")

    if potriviri:
        print("═" * 60)
        print("\nPOSIBILE POTRIVIRI — de confirmat de tine\n")
        for pot in potriviri:
            print(f"  PROIECT: {pot['proiect']['titlu'][:90]}")
            print(f"           publicat {pot['proiect']['data']}")
            for c in pot["candidati"]:
                print(f"    → {c['scor']:.0%}  {c['titlu'][:85]}")
                print(f"           ({c['motiv']}; sursa: {c['sursa']})")
            print("    Dacă e corect, în index.json pune:")
            print('      "stare": "adoptat",  "act_publicat": "<nr. și M.Of.>"\n')

    if not noi and not potriviri:
        print("Nimic de raportat. Nu se trimite nimic.")

    if AVERTISMENTE:
        print("\n" + "═" * 60)
        print(f"\nDE VERIFICAT MANUAL — {len(AVERTISMENTE)}\n")
        for a in AVERTISMENTE:
            print(f"  • {a}")
        print("\n  Nimic nu s-a pierdut. Doar nu s-a procesat automat.\n")

    if alarme:
        print("═" * 60)
        print(f"\n /!\\  ALARMĂ — {len(alarme)} surse care mergeau au picat\n")
        for a in alarme:
            print(f"  • {a}")
        print("\n  Astea funcționau. Cât timp sunt căzute, s-ar putea să ratezi acte.\n")

    if de_reparat:
        print("═" * 60)
        print(f"\nDE REPARAT — {len(de_reparat)} surse care n-au mers niciodată\n")
        for a in de_reparat:
            print(f"  • {a}")
        print("\n  Nu e urgent, dar nici nu dispare. Când găsești adresa corectă,")
        print("  o schimbi în SURSE și pui True la sfârșit.\n")


if __name__ == "__main__":
    main()
